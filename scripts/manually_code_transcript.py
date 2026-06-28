import re
import argparse
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path

'''
manually code transcript, given csv codebook
'''

COLOR_MAP = {
    "YELLOW": WD_COLOR_INDEX.YELLOW,
    "PINK": WD_COLOR_INDEX.PINK,
    "TURQUOISE": WD_COLOR_INDEX.TURQUOISE,
    "BRIGHT_GREEN": WD_COLOR_INDEX.BRIGHT_GREEN,
    "BLUE": WD_COLOR_INDEX.BLUE,
    "RED": WD_COLOR_INDEX.RED,
    "GREEN": WD_COLOR_INDEX.GREEN,
    "GRAY_25": WD_COLOR_INDEX.GRAY_25,
    "NONE": None,
    "": None,
}


# THEMES
# -----------------------------
'''
THEMES = {
    1: "AI",
    2: "Peer Relationships",
    3: "Teaching Methods",
    4: "Information Sources & Engagement",
    0: "NONE",
}

HIGHLIGHT = {
    1: WD_COLOR_INDEX.YELLOW,        # AI
    2: WD_COLOR_INDEX.PINK,          # Peer Relationshios
    3: WD_COLOR_INDEX.TURQUOISE,     # Teaching Methods
    4: WD_COLOR_INDEX.BRIGHT_GREEN,  # Information Sources
    0: None, 
}
'''

# -----------------------------
# TEXT HANDLING
# -----------------------------
def read_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_transcript(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    return read_txt(path)

def sentence_split(text):
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    return [s.strip() for s in sentences if s.strip()]


def load_codebook(path: Path) -> dict[int, dict]:
    df = pd.read_csv(path)

    required_columns = {"code", "name", "color"}
    if not required_columns.issubset(df.columns):
        raise ValueError("Codebook must contain columns: code, name, color")

    codebook = {}

    for i, row in df.iterrows():
        code = int(row["code"])
        color_name = str(row.get("color", "")).strip().upper()

        codebook[code] = {
            "name": str(row["name"]),
            "color": COLOR_MAP.get(color_name, None),
        }

    return codebook

# -----------------------------
# INTERACTIVE CODING
# -----------------------------
def print_menu(codebook: dict[int, dict]) -> None:
    print("\nChoose a code:")
    for code, info in sorted(codebook.items()):
        if code != 0:
            print(f" {code} = {info['name']}")
    print(" Enter = NONE")
    print(" b = back | q = quit\n")

def manual_code(sentences: list[str], codebook: dict[int, dict]) -> list[dict]:
    decisions = []
    i = 0

    while i < len(sentences):
        print("\n" + "=" * 90)
        print(f"Sentence {i+1}/{len(sentences)}")
        print("-" * 90)
        print(sentences[i])
        print_menu(codebook)

        cmd = input("Your choice: ").strip().lower()

        if cmd == "":
            code = 0
        elif cmd == "q":
            break
        elif cmd == "b":
            if decisions:
                decisions.pop()
                i -= 1
            else:
                print("Already at the beginning.")
            continue
        elif cmd.isdigit() and int(cmd) in codebook:
            code = int(cmd)
        else:
            print("Invalid input. Try again.")
            continue

        decisions.append(
            {
                "sentence_id": i + 1,
                "text": sentences[i],
                "theme_code": code,
                "theme_name": codebook[code]["name"],
            }
        )

        i += 1

    return decisions

# -----------------------------
# OUTPUT
# -----------------------------
def write_outputs(decisions: list[dict],
    codebook: dict[int, dict],
    out_csv: Path,
    out_docx: Path,) -> None:
    # CSV audit trail
    out_csv.parent.mkdir(parents=True, exist_ok=True)
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    df = pd.DataFrame(decisions)
    df.to_csv(out_csv, index=False)

    doc = Document()
    doc.add_heading("Highlighted Transcript", level=1)

    legend = doc.add_paragraph("Legend: ")

    for code, info in sorted(codebook.items()):
        label = f"[{code}] {info['name']}  "
        run = legend.add_run(label)

        if info["color"]:
            run.font.highlight_color = info["color"]

    doc.add_paragraph("")

    for decision in decisions:
        p = doc.add_paragraph()
        run = p.add_run(decision["text"])

        color = codebook[decision["theme_code"]]["color"]
        if color:
            run.font.highlight_color = color

    doc.save(out_docx)

# -----------------------------
# MAIN
# -----------------------------
def main() -> None:
    parser = argparse.ArgumentParser(
        description="Manually code a transcript using CSV codebook."
    )

    parser.add_argument(
        "input_file",
        help="Transcript file to code (.txt or .docx).",
    )

    parser.add_argument(
        "--codebook",
        required=True,
        help="CSV file with columns: code, name, color.",
    )

    parser.add_argument(
        "--out-dir",
        default="outputs",
        help="Directory where coded outputs will be saved.",
    )

    args = parser.parse_args()

    input_path = Path(args.input_file)
    codebook_path = Path(args.codebook)
    out_dir = Path(args.out_dir)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if not codebook_path.exists():
        raise FileNotFoundError(f"Codebook not found: {codebook_path}")

    codebook = load_codebook(codebook_path)
    raw_text = read_transcript(input_path)
    sentences = sentence_split(raw_text)

    decisions = manual_code(sentences, codebook)

    out_csv = out_dir / f"{input_path.stem}_coding_decisions.csv"
    out_docx = out_dir / f"{input_path.stem}_highlighted.docx"

    write_outputs(decisions, codebook, out_csv, out_docx)

    print("\nDone.")
    print(f"Saved: {out_csv}")
    print(f"Saved: {out_docx}")


if __name__ == "__main__":
    main()